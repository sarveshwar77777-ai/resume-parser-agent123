import io
import pdfplumber
import docx
from models import ExtractorResult, ExtractorStatus

def extract(data: bytes, filename: str) -> ExtractorResult:
    """Extracts text from PDF or DOCX file bytes."""
    raw_text = ""
    ext = filename.lower().split('.')[-1]
    
    if ext not in ["pdf", "docx"]:
        return ExtractorResult(
            raw_text="",
            status=ExtractorStatus.UNKNOWN,
            message=f"UNKNOWN - unsupported extension '{ext}'"
        )
        
    try:
        if ext == "pdf":
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text(layout=True)
                    if text and text.strip():
                        pages_text.append(text)
                if not pages_text:
                    return ExtractorResult(
                        raw_text="",
                        status=ExtractorStatus.UNKNOWN,
                        message="UNKNOWN - no text layer detected"
                    )
                raw_text = "\n\n".join(pages_text)
                
        elif ext == "docx":
            doc = docx.Document(io.BytesIO(data))
            parts = []
            for para in doc.paragraphs:
                parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            raw_text = "\n".join(parts)
            
        return ExtractorResult(
            raw_text=raw_text,
            status=ExtractorStatus.OK,
            message="Success"
        )
    except Exception as e:
        return ExtractorResult(
            raw_text="",
            status=ExtractorStatus.UNKNOWN,
            message=f"UNKNOWN - {str(e)}"
        )
